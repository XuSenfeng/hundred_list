from docx import Document
# 设置英寸像素点
from docx.shared import Inches, Pt, Cm
# 设置段落对齐
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 设置字体
from docx.oxml.ns import qn
# 设置字体颜色色
from docx.shared import RGBColor


vdoc = Document()
vdoc.add_heading("这是一个测试标题", level=9)

vparagraph = vdoc.add_paragraph("段落2： 用于替换测试的段落")
vparagraph = vdoc.add_paragraph("段落3: 段落的开始部分", style='List Bullet')

vparagraph.add_run("段落第二句话,")
vparagraph.add_run("段落第三句话。")

vparagraph2 = vdoc.add_paragraph("段落4，这是另一个段落， 这是另一个额段落， 这是另一个段落， 这是另一个段落")
run = vparagraph2.add_run("这是黑体，")
run.bold = True
run1 = vparagraph2.add_run("这是斜体，")
run1.italic = True
run2 = vparagraph2.add_run("这是下划线，")
run1.underline = True
run3 = vparagraph2.add_run("这是正方黑体")

run3.font.name = u"正方黑体"
run_el = run3._element
run_el.rPr.rFonts.set(qn('w:eastAsia'), '正方黑体')
# 设置字体的大小
run3.font.size = Pt(16)
run3.font.color.rgb = RGBColor(255, 0, 0)

paragraphs_list = vdoc.paragraphs

vformot = paragraphs_list[3].paragraph_format
vformot.first_line_indent = Inches(0.6)

# 设置段落的缩进
vformot.left_indent = Pt(10)
vformot.line_spacing = Pt(18)
vformot.space_before = Pt(16)
vformot.space_after = Pt(16)

vformot.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

paragraphs_list = vdoc.paragraphs

paragraphs_list[1].text = "这一段的字体替换了"
print(paragraphs_list[1].text)

paragraphs_4 = paragraphs_list[3]

p_paragraph = paragraphs_4.insert_paragraph_before("第四段前")
# 增加一页
vdoc.add_page_break()
# 加入一个图片
vdoc.add_picture("1.jpg", width=Cm(2))
vdoc.save("./word.docx")



