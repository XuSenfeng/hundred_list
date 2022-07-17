from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor

info = [
    {'name': '李明', 'sendtype': '电子邮件', 'offer': '企业'},
    {'name': '张三', 'sendtype': '学校推荐', 'offer': '信息管理'},
    {'name': '李四', 'sendtype': '电话沟通', 'offer': '软件开发'},
    {'name': '王五', 'sendtype': '电子邮件', 'offer': '经营管理'},
    {'name': '赵六', 'sendtype': '网络表哥', 'offer': '广告推广'},
]

for person_info in info:
    vdoc = Document()
    p0 = vdoc.add_paragraph()
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p0_format = p0.paragraph_format
    p0_format.space_before = Pt(8)
    p0_format.space_after = Pt(32)

    run0 = p0.add_run("面试通知书")

    run0.font.name = u"黑体"
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u"微软黑体")
    run0.font.size = Pt(21)
    run0.font.bold = True

    p1 = vdoc.add_paragraph()

    run1 = p1.add_run(person_info['name']+ "先生/女士您好")
    run1.font.name = '仿宋'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋")
    run1.font.size = Pt(16)
    run1.font.bold = True

    p2 = vdoc.add_paragraph()

    p2.paragraph_format.first_line_indent = Pt(32)
    run3 = p2.add_run(f"我公司通过{person_info['sendtype']}获得了您对我们公司的信任与选择， "
                      f"经过人力部的初步审查， 认为您具有{person_info['offer']}的资格， "
                      f"通知您来面试")
    run3.font.name = '仿宋'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋")
    run3.font.size = Pt(16)
    # 第三段
    p3 = vdoc.add_paragraph()

    p3.paragraph_format.first_line_indent = Pt(32)
    run4 = p3.add_run("具体要求如下")
    run4.font.name = '仿宋'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋")
    run4.font.size = Pt(16)

    # 第四段
    p4 = vdoc.add_paragraph()

    p4.paragraph_format.first_line_indent = Pt(32)
    run5 = p4.add_run("面试时间：2020年6月8日上午10点")
    run5.font.name = '仿宋'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋")
    run5.font.size = Pt(16)

    # 第五段
    p5 = vdoc.add_paragraph()

    p5.paragraph_format.first_line_indent = Pt(32)
    run4 = p5.add_run("面试地点：啥的沙发是的方式GV ")
    run4.font.name = '仿宋'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋")
    run4.font.size = Pt(16)

    vdoc.save('./notice_dir/notice_' + person_info['name']+'.docx')
